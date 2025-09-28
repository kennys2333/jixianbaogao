import sys
import re
import os
import hashlib
from datetime import datetime
from bs4 import BeautifulSoup
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QTreeWidget, QTreeWidgetItem, QTextEdit,
                            QSplitter, QLabel, QComboBox, QPushButton, QFileDialog,
                            QMessageBox, QProgressBar, QGroupBox)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor, QFontDatabase
# DOC处理库
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# -------------------------- 数据提取线程 --------------------------
class FileExtractorThread(QThread):
    progress_updated = pyqtSignal(int)
    extraction_finished = pyqtSignal(list)
    
    def __init__(self, directory):
        super().__init__()
        self.directory = directory
        self.config_items = []
        self.running = True
        
    def run(self):
        try:
            # 完全迭代方式获取所有HTML文件
            html_files = []
            stack = [self.directory]
            while stack and self.running:
                current_dir = stack.pop()
                try:
                    entries = os.listdir(current_dir)
                except Exception as e:
                    print(f"无法访问目录 {current_dir}: {e}")
                    continue
                    
                for entry in entries:
                    entry_path = os.path.join(current_dir, entry)
                    if os.path.isdir(entry_path):
                        stack.append(entry_path)
                    elif os.path.isfile(entry_path) and entry.lower().endswith(('.html', '.htm')):
                        html_files.append(entry_path)
            
            if not html_files:
                self.extraction_finished.emit([])
                return
                
            total_files = len(html_files)
            for i in range(total_files):
                if not self.running:
                    break
                    
                file_path = html_files[i]
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                    
                    self.extract_from_html(html_content, file_path)
                    self.progress_updated.emit(int((i + 1) / total_files * 100))
                except Exception as e:
                    print(f"处理文件 {file_path} 时出错: {str(e)}")
            
            self.extraction_finished.emit(self.config_items)
        except Exception as e:
            print(f"提取过程出错: {str(e)}")
            self.extraction_finished.emit([])
    
    def extract_from_html(self, html_content, file_path):
        ip_address = self.extract_ip_from_filename(os.path.basename(file_path))
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 查找Linux配置规范部分
        linux_config_section = None
        h3_tags = soup.find_all(class_='report_h3')
        for h3 in h3_tags:
            if 'Linux 配置规范' in h3.text:
                linux_config_section = h3
                break
        
        if not linux_config_section:
            return
        
        config_table = linux_config_section.find_next('table', id='config_info')
        if not config_table:
            return
        
        current_category = None
        rows = config_table.find_all('tr')
        
        # 迭代处理所有行
        for row in rows:
            row_classes = row.get('class', [])
            
            if 'second_title' in row_classes:
                td = row.find('td')
                if td:
                    current_category = td.text.strip()
                continue
            
            if 'first_title' in row_classes:
                continue
            
            if 'baseline' in row_classes and any(cls in row_classes for cls in ['baseline_yes', 'baseline_no']):
                risk_level_span = row.find('span', class_=re.compile(r'level_danger'))
                risk_level = risk_level_span['class'][0] if risk_level_span else '未知风险等级'
                
                # 转换风险等级为中文
                risk_map = {'level_danger_high': '高危', 'level_danger_middle': '中危', 'level_danger_low': '低危'}
                risk_text = risk_map.get(risk_level, '未知')
                
                if 'low' in risk_level:
                    continue
                
                description = risk_level_span.text.strip() if risk_level_span else row.find('td').text.strip()
                baseline_id = None
                for cls in row_classes:
                    if cls.startswith('baseline_id_'):
                        baseline_id = cls
                        break
                
                checkpoints = []
                config_methods = {}
                if baseline_id:
                    # 收集所有相关详细行
                    detail_rows = []
                    for r in rows:
                        if baseline_id in r.get('class', []) and r != row:
                            detail_rows.append(r)
                    
                    # 处理详细行
                    for detail_row in detail_rows:
                        detail_table = detail_row.find('table', class_='report_table plumb')
                        if detail_table:
                            data_rows = detail_table.find_all('tr')[1:]  # 跳过表头
                            seen_checkpoints = set()
                            row_index = 0
                            
                            for data_row in data_rows:
                                row_index += 1
                                # 只处理error.gif的项
                                result_img = data_row.find('img', src=re.compile(r'is_safe\.gif', re.IGNORECASE))
                                if result_img:
                                    continue
                                
                                error_img = data_row.find('img', src=re.compile(r'error\.gif', re.IGNORECASE))
                                if not error_img:
                                    continue
                                
                                check_result = "不合格"
                                tds = data_row.find_all('td')
                                if len(tds) < 6:
                                    continue
                                
                                # 修复列偏移逻辑：根据tds数量确定列索引
                                if len(tds) == 7:
                                    cols = [1, 2, 3, 4, 5]  # checkpoint, actual, match, standard, method
                                else:
                                    cols = [0, 1, 2, 3, 4]
                                
                                checkpoint_name = tds[cols[0]].text.strip()
                                actual_value = tds[cols[1]].text.strip()
                                match_rule = tds[cols[2]].text.strip()
                                standard_value = tds[cols[3]].text.strip()
                                config_method = self.clean_config_method(tds[cols[4]].text.strip())
                                
                                if not checkpoint_name:
                                    continue
                                
                                # 创建检查点唯一标识，用于去重
                                checkpoint_key = f"{checkpoint_name}_{actual_value}_{match_rule}_{standard_value}"
                                
                                # 只添加未出现过的检查点
                                if checkpoint_key not in seen_checkpoints:
                                    seen_checkpoints.add(checkpoint_key)
                                    
                                    config_hash = hashlib.md5(config_method.encode()).hexdigest()[:8]
                                    config_methods[config_hash] = config_method
                                    
                                    checkpoints.append({
                                        'checkpoint': checkpoint_name,
                                        'actual_value': actual_value,
                                        'match_rule': match_rule,
                                        'standard_value': standard_value,
                                        'config_hash': config_hash,
                                        'check_result': check_result,
                                        'original_order': row_index
                                    })
                
                if checkpoints:
                    # 按原始顺序排序
                    checkpoints.sort(key=lambda x: x['original_order'])
                    self.config_items.append({
                        'risk_level': risk_text,
                        'description': description,
                        'category': current_category or '未分类',
                        'source_file': os.path.basename(file_path),
                        'ip_address': ip_address,
                        'checkpoints': checkpoints,
                        'config_methods': config_methods,
                        'check_type': "系统维护"
                    })
    
    def clean_config_method(self, config_text):
        lines = [line.strip() for line in config_text.split('\n') if line.strip()]
        cleaned = '; '.join(lines).replace(',;', ';').replace(';;', ';')
        return cleaned.replace('; ', '\n')
    
    def extract_ip_from_filename(self, filename):
        ip_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
        match = re.search(ip_pattern, filename)
        return match.group() if match else filename
    
    def stop(self):
        self.running = False
        self.wait()

# -------------------------- GUI主窗口 --------------------------
class ConfigurationExtractorGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.merged_items = []
        self.category_hierarchy = {}
        self.initUI()
        self.current_directory = ""
        self.extractor_thread = None
        
    def initUI(self):
        self.setWindowTitle('Linux配置规范提取工具（kennys-2025-09-23）')
        self.setGeometry(100, 100, 1400, 900)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 顶部控制区域
        control_layout = QHBoxLayout()
        self.dir_label = QLabel('未选择目录')
        self.dir_label.setStyleSheet("QLabel { background-color: #333333; padding: 5px; }")
        self.select_dir_btn = QPushButton('选择目录')
        self.select_dir_btn.clicked.connect(self.select_directory)
        self.load_btn = QPushButton('加载所有HTML文件')
        self.load_btn.clicked.connect(self.load_all_html_files)
        self.load_btn.setEnabled(False)
        
        self.export_btn = QPushButton('导出为DOCX')
        self.export_btn.clicked.connect(self.export_to_docx)
        self.export_btn.setEnabled(False)
        
        self.filter_combo = QComboBox()
        self.filter_combo.addItems(['所有风险等级(不含低风险)', '高危', '中危'])
        self.filter_combo.currentTextChanged.connect(self.filter_items)
        
        control_layout.addWidget(self.dir_label, 3)
        control_layout.addWidget(self.select_dir_btn, 1)
        control_layout.addWidget(self.load_btn, 1)
        control_layout.addWidget(self.export_btn, 1)
        control_layout.addWidget(QLabel('风险过滤:'))
        control_layout.addWidget(self.filter_combo, 1)
        main_layout.addLayout(control_layout)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        # 统计信息
        self.stats_group = QGroupBox("统计信息")
        stats_layout = QHBoxLayout()
        self.file_count_label = QLabel("处理文件: 0")
        self.item_count_label = QLabel("提取检查项: 0")
        self.unique_issues_label = QLabel("唯一问题: 0")
        stats_layout.addWidget(self.file_count_label)
        stats_layout.addWidget(self.item_count_label)
        stats_layout.addWidget(self.unique_issues_label)
        self.stats_group.setLayout(stats_layout)
        main_layout.addWidget(self.stats_group)
        
        # 分割器（左右布局）
        splitter = QSplitter(Qt.Horizontal)
        self.tree_widget = QTreeWidget()
        self.tree_widget.setHeaderLabels(['检查项', '风险等级', '受影响IP数量'])
        self.tree_widget.setColumnWidth(0, 500)
        self.tree_widget.setColumnWidth(1, 100)
        self.tree_widget.setColumnWidth(2, 120)
        self.tree_widget.itemSelectionChanged.connect(self.show_details)
        splitter.addWidget(self.tree_widget)
        
        self.detail_text = QTextEdit()
        self.detail_text.setReadOnly(True)
        # 正确实例化QFontDatabase并获取字体列表
        font_db = QFontDatabase()
        font_families = font_db.families()
        
        # 选择合适的等宽字体
        if "Consolas" in font_families:
            fixed_font = QFont("Consolas")
        elif "Monaco" in font_families:
            fixed_font = QFont("Monaco")  # macOS
        elif "Courier New" in font_families:
            fixed_font = QFont("Courier New")  # Windows
        elif "Noto Mono" in font_families:
            fixed_font = QFont("Noto Mono")  # Linux
        else:
            # 最后尝试使用系统默认等宽字体
            fixed_font = QFont()
            fixed_font.setStyleHint(QFont.Monospace)
        
        self.detail_text.setFont(fixed_font)
        self.detail_text.setLineWrapMode(QTextEdit.NoWrap)
        splitter.addWidget(self.detail_text)
        splitter.setSizes([700, 700])
        main_layout.addWidget(splitter, 1)
        
        self.statusBar().showMessage('就绪')
        
    def select_directory(self):
        directory = QFileDialog.getExistingDirectory(
            self, '选择目录', '', QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
        )
        if directory:
            self.current_directory = directory
            self.dir_label.setText(f'目录: {os.path.basename(directory)}')
            self.load_btn.setEnabled(True)
            self.statusBar().showMessage(f'已选择目录: {directory}')
            
            # 迭代计算HTML文件数量
            html_count = 0
            stack = [directory]
            while stack:
                current_dir = stack.pop()
                try:
                    entries = os.listdir(current_dir)
                except Exception as e:
                    print(f"无法访问目录 {current_dir}: {e}")
                    continue
                    
                for entry in entries:
                    entry_path = os.path.join(current_dir, entry)
                    if os.path.isdir(entry_path):
                        stack.append(entry_path)
                    elif os.path.isfile(entry_path) and entry.lower().endswith(('.html', '.htm')):
                        html_count += 1
            
            self.file_count_label.setText(f"处理文件: {html_count}")
            self.item_count_label.setText(f"提取检查项: 0")
            self.unique_issues_label.setText(f"唯一问题: 0")
    
    def load_all_html_files(self):
        if not self.current_directory:
            QMessageBox.warning(self, '警告', '请先选择目录')
            return
            
        self.tree_widget.clear()
        self.detail_text.clear()
        self.statusBar().showMessage('正在提取配置信息...')
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.load_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        
        self.extractor_thread = FileExtractorThread(self.current_directory)
        self.extractor_thread.progress_updated.connect(self.update_progress)
        self.extractor_thread.extraction_finished.connect(self.on_extraction_finished)
        self.extractor_thread.start()
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
        self.statusBar().showMessage(f'正在提取配置信息... {value}%')
    
    def on_extraction_finished(self, config_items):
        self.progress_bar.setVisible(False)
        self.load_btn.setEnabled(True)
        
        if not config_items:
            QMessageBox.information(self, '提示', '未提取到任何不合格的配置信息')
            self.statusBar().showMessage('提取完成，未找到不合格项')
            return
        
        # 合并相同问题
        self.merged_items = self.merge_same_issues(config_items)
        self.export_btn.setEnabled(True)
        
        # 构建分类层级（只保留主类）
        self.build_main_category_hierarchy()
        
        # 更新统计信息
        unique_files = set(item['source_file'] for item in config_items)
        self.file_count_label.setText(f"处理文件: {len(unique_files)}")
        self.item_count_label.setText(f"提取检查项: {len(config_items)}")
        self.unique_issues_label.setText(f"唯一问题: {len(self.merged_items)}")
        
        # 构建树形视图
        self.build_tree_view()
        
        self.tree_widget.expandAll()
        self.statusBar().showMessage(f'提取完成，共找到 {len(self.merged_items)} 个唯一问题')
    
    def build_main_category_hierarchy(self):
        """只构建主分类层级，不保留子分类"""
        self.category_hierarchy = {}
        
        # 迭代处理所有合并项，只保留主分类
        for item in self.merged_items:
            main_category = self.get_main_category(item['category'])
            
            if main_category not in self.category_hierarchy:
                self.category_hierarchy[main_category] = {
                    'id': len(self.category_hierarchy) + 1,
                    'items': []  # 直接存放项目，不设子分类
                }
            
            self.category_hierarchy[main_category]['items'].append(item)
    
    def get_main_category(self, category):
        """提取主分类，忽略子分类"""
        if '认证' in category:
            return '认证授权'
        elif '协议' in category:
            return '协议安全'
        elif '密码' in category:
            return '密码策略'
        elif 'FTP' in category:
            return '服务安全'
        elif '日志' in category:
            return '日志审计'
        elif '防火墙' in category:
            return '防火墙配置'
        else:
            # 取第一个空格前的部分作为主分类
            return category.split(' ')[0] if category else '其他'
    
    def build_tree_view(self):
        # 清空现有项
        self.tree_widget.clear()
        
        # 只按主分类添加项目
        main_categories = list(self.category_hierarchy.items())
        for cat_name, cat_data in main_categories:
            cat_item = QTreeWidgetItem(self.tree_widget, [cat_name, '', ''])
            cat_item.setExpanded(True)
            
            # 直接添加项目，不分级
            for item in cat_data['items']:
                risk_text = item['risk_level']
                ip_count = len(item['ip_addresses'])
                item_widget = QTreeWidgetItem(cat_item, [item['description'], risk_text, f"{ip_count}个"])
                
                # 风险等级着色
                color_map = {'高危': QColor('red'), '中危': QColor('orange')}
                if risk_text in color_map:
                    item_widget.setForeground(1, color_map[risk_text])
                
                item_widget.setData(0, Qt.UserRole, item)
    
    def merge_same_issues(self, config_items):
        merged = {}
        for item in config_items:
            # 使用检查点的原始顺序来生成唯一键，确保顺序一致
            checkpoint_keys = []
            for cp in item['checkpoints']:
                checkpoint_keys.append(f"{cp['checkpoint']}_{cp['config_hash']}_{cp['original_order']}")
            
            unique_key = f"{item['description']}_{'_'.join(checkpoint_keys)}"
            
            if unique_key not in merged:
                merged_item = {
                    'risk_level': item['risk_level'],
                    'description': item['description'],
                    'category': item['category'],
                    'ip_addresses': [item['ip_address']],
                    'source_files': [item['source_file']],
                    'checkpoints': item['checkpoints'],
                    'config_methods': item['config_methods'],
                    'check_type': item['check_type']
                }
                merged[unique_key] = merged_item
            else:
                if item['ip_address'] not in merged[unique_key]['ip_addresses']:
                    merged[unique_key]['ip_addresses'].append(item['ip_address'])
        
        return list(merged.values())
    
    def show_details(self):
        current_item = self.tree_widget.currentItem()
        if current_item and current_item.childCount() == 0:
            item_data = current_item.data(0, Qt.UserRole)
            if item_data:
                details = f"风险等级: {item_data['risk_level']}\n"
                details += f"检查类型: {item_data['check_type']}\n"
                details += f"描述: {item_data['description']}\n"
                details += f"受影响IP地址: {', '.join(item_data['ip_addresses'])}\n"
                details += f"受影响数量: {len(item_data['ip_addresses'])}\n\n"
                
                if item_data.get('checkpoints') and item_data.get('config_methods'):
                    details += "检查点详情:\n"
                    details += "-" * 180 + "\n"
                    for i, checkpoint in enumerate(item_data['checkpoints'], 1):
                        details += f"{i}. {checkpoint['checkpoint']}（{checkpoint['check_result']}）\n"
                        details += f"   验证结果: 实际值：{checkpoint['actual_value']:10} "
                        details += f"匹配规则: {checkpoint['match_rule']:8} "
                        details += f"标准值: {checkpoint['standard_value']}\n"
                        details += "-" * 180 + "\n"
                    
                    details += "修复建议:\n"
                    details += "-" * 180 + "\n"
                    config_hashes = set(cp['config_hash'] for cp in item_data['checkpoints'])
                    unique_configs = {h: item_data['config_methods'][h] for h in config_hashes}
                    for config in unique_configs.values():
                        details += f"{config}\n"
                        details += "-" * 180 + "\n"
                
                self.detail_text.setText(details)
    
    def filter_items(self, filter_text):
        # 迭代过滤所有项目
        for top_idx in range(self.tree_widget.topLevelItemCount()):
            top_item = self.tree_widget.topLevelItem(top_idx)
            for item_idx in range(top_item.childCount()):
                item = top_item.child(item_idx)
                item_data = item.data(0, Qt.UserRole)
                
                if not item_data:
                    continue
                    
                if filter_text == '所有风险等级(不含低风险)':
                    item.setHidden(False)
                elif filter_text == '高危' and item_data['risk_level'] == '高危':
                    item.setHidden(False)
                elif filter_text == '中危' and item_data['risk_level'] == '中危':
                    item.setHidden(False)
                else:
                    item.setHidden(True)
    
    def set_cell_border(self, cell):
        """使用parse_xml创建正确的元素类型"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 检查是否已有边框设置
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            # 使用XML字符串创建正确的元素类型
            tcBorders = parse_xml('''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                </w:tcBorders>
            ''')
            tcPr.append(tcBorders)
    
    def set_table_wrap_around(self, table):
        """设置表格文字环绕"""
        tbl_pr = table._tbl.tblPr
        
        # 使用parse_xml创建正确的元素类型
        tbl_wrap = parse_xml('''
            <w:tblWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="around"/>
        ''')
        tbl_pr.append(tbl_wrap)
        
        tbl_pos = parse_xml('''
            <w:tblPosition xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="left"/>
        ''')
        tbl_pr.append(tbl_pos)
    
    def export_to_docx(self):
        """导出为指定格式的表格：三级编号，主分类显示，黑色文字"""
        if not self.merged_items:
            QMessageBox.warning(self, '警告', '没有可导出的数据，请先加载文件')
            return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, "导出为DOCX", os.path.expanduser("~"), "Word文档 (*.docx)"
        )
        if not save_path:
            return
        
        try:
            doc = Document()
            
            # 设置全局字体为宋体小四号，黑色
            style = doc.styles['Normal']
            style.font.name = u'宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 迭代处理所有主分类和项目
            main_categories = list(self.category_hierarchy.items())
            for cat_idx, (cat_name, cat_data) in enumerate(main_categories, 1):
                # 主分类标题（一级编号）
                cat_heading = doc.add_heading(f'1.{cat_idx}.{cat_name}', level=1)
                self.set_heading_style(cat_heading, Pt(14))
                
                # 检查项（二级编号）
                for item_idx, item in enumerate(cat_data['items'], 1):
                    # 检查项标题（三级编号：1.1.1.标题）
                    item_number = f'1.{cat_idx}.{item_idx}'
                    # 使用描述作为标题，不拼接检查点
                    item_title = f'{item_number}.{item["description"]}（不合格）'
                    
                    item_heading = doc.add_heading(item_title, level=2)
                    self.set_heading_style(item_heading, Pt(12))
                    
                    # 创建6行2列的表格
                    table = doc.add_table(rows=6, cols=2)
                    
                    # 表格总宽度15.24厘米
                    table.width = Cm(15.24)
                    
                    # 左对齐
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                    # 文字环绕
                    self.set_table_wrap_around(table)
                    
                    # 列宽设置：第二列12.25厘米，第一列2.99厘米
                    table.columns[0].width = Cm(2.99)
                    table.columns[1].width = Cm(12.25)
                    
                    # 填充表格内容
                    rows = table.rows
                    rows[0].cells[0].text = "风险等级"
                    rows[0].cells[1].text = item['risk_level']
                    
                    rows[1].cells[0].text = "检查类型"
                    rows[1].cells[1].text = "系统维护"
                    
                    rows[2].cells[0].text = "不合格主机"
                    rows[2].cells[1].text = ', '.join(item['ip_addresses'])
                    
                    rows[3].cells[0].text = "检查项描述"
                    rows[3].cells[1].text = item['description']
                    
                    rows[4].cells[0].text = "验证结果"
                    # 合并所有检查点的验证结果（去重处理）
                    verify_results = []
                    seen_verify = set()  # 用于验证结果去重
                    for checkpoint in item['checkpoints']:
                        # 创建验证结果唯一标识
                        verify_key = f"{checkpoint['actual_value']}_{checkpoint['match_rule']}_{checkpoint['standard_value']}"
                        if verify_key not in seen_verify:
                            seen_verify.add(verify_key)
                            verify_results.append(
                                f"实际值：{checkpoint['actual_value']} "
                                f"匹配规则：{checkpoint['match_rule']} "
                                f"标准值：{checkpoint['standard_value']}"
                            )
                    # 将分号替换为换行符
                    rows[4].cells[1].text = '\n'.join(verify_results)
                    
                    rows[5].cells[0].text = "修复建议"
                    # 合并所有检查点的修复建议
                    config_hashes = set(cp['config_hash'] for cp in item['checkpoints'])
                    unique_configs = {h: item['config_methods'][h] for h in config_hashes}
                    rows[5].cells[1].text = '\n\n'.join(unique_configs.values())
                    
                    # 设置边框和字体
                    for row in rows:
                        for cell in row.cells:
                            self.set_cell_border(cell)
                            for para in cell.paragraphs:
                                para.style = style
                                # 验证结果左对齐，方便阅读
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 添加空行分隔
                    doc.add_paragraph('')
            
            # 保存文档
            doc.save(save_path)
            QMessageBox.information(self, '成功', f'文档已导出至:\n{save_path}')
            self.statusBar().showMessage(f'文档导出成功: {save_path}')
            
        except Exception as e:
            QMessageBox.critical(self, '导出失败', f'导出过程出错: {str(e)}')
            print(f"导出错误: {str(e)}")
    
    def set_heading_style(self, heading, font_size):
        """设置标题样式为黑色"""
        for run in heading.runs:
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = font_size
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 确保文字为黑色
    
    def closeEvent(self, event):
        if self.extractor_thread and self.extractor_thread.isRunning():
            self.extractor_thread.stop()
        event.accept()

# -------------------------- 程序入口 --------------------------
def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = ConfigurationExtractorGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
    
